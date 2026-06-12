from io import BytesIO

from docx import Document
from docx.shared import Pt


def add_heading(document, text, level=1):
    document.add_heading(str(text or ""), level=level)


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(str(text or ""))
    run.font.size = Pt(10)


def build_word_report(report: dict) -> bytes:
    document = Document()

    document.add_heading("E-book / PDF Marketing QA Report", level=0)

    add_heading(document, "1. Executive Summary", level=1)
    add_paragraph(document, f"Overall QA status: {report.get('overall_status', '')}")
    add_paragraph(document, f"Overall quality score: {report.get('quality_score', '')}/100")
    add_paragraph(document, report.get("executive_summary", ""))

    add_heading(document, "Top 5 Recommended Fixes", level=2)
    for item in report.get("top_recommendations", [])[:5]:
        document.add_paragraph(str(item), style="List Bullet")

    add_heading(document, "2. Issue Log", level=1)

    issues = report.get("issues", [])

    if issues:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        headers = [
            "Page/Section",
            "Category",
            "Severity",
            "Issue",
            "Explanation",
            "Recommended Fix"
        ]

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for issue in issues:
            row = table.add_row().cells
            row[0].text = str(issue.get("page_or_section", ""))
            row[1].text = str(issue.get("category", ""))
            row[2].text = str(issue.get("severity", ""))
            row[3].text = str(issue.get("issue", ""))
            row[4].text = str(issue.get("explanation", ""))
            row[5].text = str(issue.get("recommended_fix", ""))
    else:
        add_paragraph(document, "No issues were identified.")

    add_heading(document, "3. Category-by-Category Review", level=1)

    for category in report.get("category_reviews", []):
        add_heading(document, category.get("category", ""), level=2)
        add_paragraph(document, f"Status: {category.get('status', '')}")
        add_paragraph(document, f"Notes: {category.get('notes', '')}")

        examples = category.get("examples", [])
        if examples:
            add_paragraph(document, "Key examples:")
            for example in examples:
                document.add_paragraph(str(example), style="List Bullet")

    add_heading(document, "4. Marketing Effectiveness Review", level=1)

    marketing = report.get("marketing_effectiveness", {})

    add_paragraph(document, f"Trust-building: {marketing.get('trust_building', '')}")
    add_paragraph(document, f"Educational value: {marketing.get('educational_value', '')}")
    add_paragraph(document, f"Brand alignment: {marketing.get('brand_alignment', '')}")
    add_paragraph(document, f"CTA quality: {marketing.get('cta_quality', '')}")
    add_paragraph(document, f"Lead-generation suitability: {marketing.get('lead_generation_suitability', '')}")

    add_heading(document, "5. Accessibility Review", level=1)
    add_paragraph(document, report.get("accessibility_review", ""))

    add_heading(document, "6. Final Recommendation", level=1)
    add_paragraph(document, report.get("final_recommendation", ""))

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream.getvalue()
